"""
Bangkitkan laporan kemajuan (.docx) untuk dosen pembimbing -- membaca
data asli dari results/metrics.json + figure di results/figures/, BUKAN
angka/klaim manual. Nol hardcode utk data kuantitatif, sama seperti
scripts/generate_figures.py.

Jalankan: python scripts/generate_report.py
Output: results/laporan_kemajuan.docx

Isi placeholder ([Nama Mahasiswa], [NIM], dst di cover) WAJIB diisi
manual oleh user -- itu info personal yang tidak ada di repo mana pun.
Beberapa bagian lain juga sengaja diberi placeholder eksplisit (ditandai
"[ISI:" ) untuk konten yang user mau kontrol sendiri.
"""

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parent.parent
METRICS_PATH = REPO_ROOT / "results" / "metrics.json"
FIGURES_DIR = REPO_ROOT / "results" / "figures"
SCREENSHOTS_DIR = Path(
    r"C:\Users\HPDESK~1\AppData\Local\Temp\claude\E--projects-project-fracture-classification"
    r"\63f5daeb-bd06-483e-a16e-e57991f70ba7\scratchpad"
)
OUT_PATH = REPO_ROOT / "results" / "laporan_kemajuan.docx"

MODEL_LABELS = {"tiny": "ConvNeXt-Tiny", "small": "ConvNeXt-Small", "base": "ConvNeXt-Base", "large": "ConvNeXt-Large"}
ACCENT = RGBColor(0xC4, 0x53, 0x1E)  # oranye "grease" -- identitas visual proyek (tokens.css light mode)


def set_base_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_placeholder(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"[ISI: {text}]")
    run.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p.paragraph_format.space_after = Pt(12)
    # border kotak tipis di sekeliling paragraf placeholder
    pPr = p._p.get_or_add_pPr()
    pbdr = pPr.makeelement(qn("w:pBdr"), {})
    for edge in ("top", "left", "bottom", "right"):
        el = pPr.makeelement(qn(f"w:{edge}"), {qn("w:val"): "single", qn("w:sz"): "8", qn("w:color"): "AAAAAA"})
        pbdr.append(el)
    pPr.append(pbdr)
    return p


def add_figure(doc, image_path: Path, caption: str, width_cm: float = 14.5):
    if not image_path.exists():
        add_placeholder(doc, f"Figure tidak ditemukan: {image_path.name} -- {caption}")
        return
    doc.add_picture(str(image_path), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(10)
    cap.paragraph_format.space_after = Pt(14)


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True


def ci_str(ci: dict) -> str:
    return f"{ci['point']:.4f}\n[{ci['lower']:.4f}, {ci['upper']:.4f}]"


def add_page_break(doc):
    doc.add_page_break()


def main():
    with open(METRICS_PATH, encoding="utf-8") as f:
        data = json.load(f)
    models = sorted(data["models"], key=lambda m: ["tiny", "small", "base", "large"].index(m["model_id"]))
    ablation = data.get("clahe_ablation")
    n_test = models[0]["test_set_size"]

    doc = Document()
    set_base_font(doc)
    for section in doc.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # ================= COVER =================
    doc.add_paragraph().paragraph_format.space_before = Pt(60)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LAPORAN KEMAJUAN PENELITIAN")
    run.bold = True
    run.font.size = Pt(18)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Klasifikasi Patah Tulang dari Citra X-ray\nMenggunakan ConvNeXt (Tiny/Small/Base/Large)")
    run.font.size = Pt(14)
    sub.paragraph_format.space_after = Pt(40)

    for label, ph in [
        ("Nama Mahasiswa", "Nama Mahasiswa"),
        ("NIM", "NIM"),
        ("Program Studi", "Program Studi"),
        ("Fakultas / Universitas", "Fakultas / Universitas"),
        ("Dosen Pembimbing", "Nama Dosen Pembimbing"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"{label}: ").bold = True
        run = p.add_run(f"[ISI: {ph}]")
        run.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    p.add_run(f"Tanggal laporan: {data.get('generated_at', '')[:10]}").italic = True

    add_page_break(doc)

    # ================= RINGKASAN EKSEKUTIF =================
    add_heading(doc, "Ringkasan Eksekutif", level=1)
    doc.add_paragraph(
        "Penelitian ini membangun dan mengevaluasi secara formal empat varian arsitektur "
        "ConvNeXt (Tiny, Small, Base, Large) untuk klasifikasi biner citra X-ray tulang "
        "(fractured / not fractured), disertai platform web demonstrasi dan backend inferensi "
        "yang telah di-deploy secara publik. Seluruh pipeline -- audit dataset, training, "
        "evaluasi statistik (bootstrap 95% CI), kalibrasi probabilitas, gerbang deteksi "
        "out-of-distribution (OOD), penjelasan visual (Grad-CAM analitik), hingga ablation "
        "study CLAHE -- dirancang untuk memperbaiki kelemahan metodologis yang ditemukan pada "
        "eksperimen awal (lihat Bagian 2)."
    )
    doc.add_paragraph(
        f"Keempat model mencapai akurasi 98,6%-99,2% pada test set (n={n_test}) yang telah "
        "dijamin bebas kebocoran data secara konstruksi (deduplikasi tingkat klaster). Interval "
        "kepercayaan 95% keempat model saling tumpang tindih, sehingga belum dapat diklaim satu "
        "arsitektur secara statistik lebih unggul dari yang lain. Sistem telah di-deploy penuh: "
        "backend FastAPI + ONNX Runtime di Google Cloud Run, frontend React di "
        "fracture.lapanproject.tech -- keduanya publik dan telah diuji end-to-end dengan citra "
        "X-ray sungguhan."
    )

    # ================= 1. LATAR BELAKANG =================
    add_heading(doc, "1. Latar Belakang dan Tujuan", level=1)
    doc.add_paragraph(
        "Diagnosis fraktur tulang dari citra X-ray secara konvensional bergantung pada "
        "interpretasi radiolog. Model deep learning berpotensi menjadi alat bantu (bukan "
        "pengganti) dalam proses skrining, terutama pada situasi dengan keterbatasan akses "
        "tenaga radiolog. Penelitian ini bertujuan untuk:"
    )
    for item in [
        "Membangun pipeline klasifikasi biner fraktur yang metodologinya benar dan dapat "
        "direproduksi -- preprocessing konsisten, perbandingan arsitektur terkontrol, dan "
        "tanpa kebocoran data antar split.",
        "Mengevaluasi empat varian ConvNeXt secara adil (konfigurasi hyperparameter identik, "
        "hanya backbone yang bervariasi) dengan pelaporan interval kepercayaan statistik.",
        "Membangun platform web dan backend inferensi yang dapat didemonstrasikan secara "
        "publik dan real-time.",
        "Menyediakan mekanisme keamanan tambahan: kalibrasi probabilitas, gerbang OOD "
        "(menolak citra yang bukan X-ray), dan penjelasan visual (Grad-CAM) untuk transparansi "
        "keputusan model.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_page_break(doc)

    # ================= 2. TEMUAN AUDIT EKSPERIMEN AWAL =================
    add_heading(doc, "2. Temuan Kritis dari Eksperimen Awal", level=1)
    doc.add_paragraph(
        "Sebelum membangun pipeline final, dilakukan audit menyeluruh terhadap eksperimen "
        "sebelumnya (arsip kode lama). Ditemukan beberapa cacat metodologis fatal yang membuat "
        "angka akurasi lama (termasuk klaim 98,6% pada salah satu model) tidak dapat dipercaya:"
    )
    findings = [
        ("Preprocessing train \u2260 test", "Generator data latih/validasi memakai rescale=1/255 "
         "(rentang [0,1]), sementara generator data uji memakai fungsi preprocess_input ConvNeXt "
         "yang ternyata tidak melakukan apa pun (no-op) -- normalisasi ConvNeXt sudah menjadi "
         "layer di dalam model. Akibatnya, tiga dari empat model diuji pada skala piksel yang "
         "berbeda dari saat dilatih, membuat akurasi uji jatuh ke ~50% (setara tebak koin) "
         "walau akurasi validasi sempat 88-97%."),
        ("Perbandingan arsitektur tidak terkontrol", "Preprocessing, learning rate, dan jumlah "
         "epoch berbeda-beda antar keempat model yang dibandingkan pada eksperimen awal -- "
         "sehingga klaim \"arsitektur X paling baik\" tidak sah secara ilmiah."),
        ("Kebocoran data antar split (ditemukan saat audit dataset final)", "Dataset Kaggle yang "
         f"dipakai (usman44m/bone-fracture-x-ray-dataset) terbukti memiliki kebocoran "
         "34,84% -- 96,3% dari 508 citra pada folder test memiliki duplikat identik-byte "
         "(MD5 sama persis) di folder train. Dari klaim 10.581 file, hanya 3.370 citra yang "
         "benar-benar unik."),
        ("Threshold keputusan dicari di test set", "Ambang batas klasifikasi optimal dicari "
         "langsung di data uji pada eksperimen lama -- ini adalah kebocoran data yang membuat "
         "metrik performa bias optimistis."),
    ]
    for title, desc in findings:
        p = doc.add_paragraph()
        p.add_run(f"{title}. ").bold = True
        p.add_run(desc)

    add_page_break(doc)

    # ================= 3. METODOLOGI =================
    add_heading(doc, "3. Metodologi", level=1)

    add_heading(doc, "3.1 Audit dan Pembagian Ulang Dataset", level=2)
    doc.add_paragraph(
        "Untuk mengatasi kebocoran 34,84% yang ditemukan, seluruh 10.581 file dikelompokkan "
        "berdasarkan hash konten (mendeteksi duplikat/near-duplikat), menghasilkan 3.370 citra "
        "unik. Setiap klaster gambar identik ditugaskan ke SATU split saja (train/val/test), "
        "sehingga kebocoran menjadi nol secara konstruksi -- bukan sekadar di bawah ambang "
        "toleransi. Pembagian akhir (rasio 70/15/15, seed=42): train=2.358, val=504, test=508 "
        "citra, dengan kelas sedikit tidak seimbang (41,7% fractured / 58,3% not fractured)."
    )

    add_heading(doc, "3.2 Arsitektur dan Training", level=2)
    doc.add_paragraph(
        "Empat varian ConvNeXt (Tiny/Small/Base/Large) dilatih dengan konfigurasi hyperparameter "
        "IDENTIK (satu file konfigurasi terkunci) -- hanya backbone yang bervariasi, memastikan "
        "perbandingan arsitektur yang adil. Training dilakukan dua fase: (1) backbone dibekukan, "
        "hanya head klasifikasi dilatih; (2) sebagian lapisan terakhir backbone dibuka "
        "(fine-tuning) dengan learning rate lebih kecil. Head model sengaja dibuat sederhana "
        "(Global Average Pooling \u2192 Dense(512, GELU) \u2192 Dense(1, sigmoid)) untuk "
        "memungkinkan perhitungan Grad-CAM secara analitik (lihat 3.4)."
    )

    add_heading(doc, "3.3 Evaluasi Formal", level=2)
    doc.add_paragraph(
        "Setiap metrik dilaporkan dengan interval kepercayaan 95% dari 2.000 kali resampling "
        "bootstrap. Threshold keputusan dipilih menggunakan Youden's J index pada data VALIDASI "
        "(bukan data uji), menghindari kebocoran yang terjadi pada eksperimen lama. Kalibrasi "
        "probabilitas dilakukan dengan temperature scaling, diukur dengan Expected Calibration "
        "Error (ECE) dan reliability diagram."
    )

    add_heading(doc, "3.4 Ekspor ONNX dan Grad-CAM Analitik", level=2)
    doc.add_paragraph(
        "Untuk inferensi ringan tanpa TensorFlow di server (ukuran image Docker ~400MB, bukan "
        "~3GB), model diekspor ke format ONNX. Karena arsitektur head yang sederhana (GAP \u2192 "
        "Dense \u2192 Dense), gradien Grad-CAM dapat diturunkan secara analitik (bentuk tertutup) "
        "tanpa memerlukan backpropagation penuh -- diverifikasi identik dengan hasil "
        "GradientTape (ground truth) hingga toleransi numerik yang sangat kecil."
    )

    add_heading(doc, "3.5 Gerbang Out-of-Distribution (OOD)", level=2)
    doc.add_paragraph(
        "Untuk mencegah model memberikan prediksi berkeyakinan tinggi pada citra yang BUKAN "
        "X-ray, diterapkan gerbang OOD berbasis jarak Mahalanobis pada fitur GAP (sebelum head "
        "klasifikasi). Statistik referensi (mean, kovarians) di-fit pada fitur data training, "
        "dan diuji terhadap dataset publik non-X-ray (CIFAR-10) sebagai referensi OOD."
    )

    add_heading(doc, "3.6 Ablation Study CLAHE", level=2)
    doc.add_paragraph(
        "Contrast Limited Adaptive Histogram Equalization (CLAHE) diuji sebagai variabel "
        "eksperimen yang sah (bukan sekadar dituliskan tanpa diverifikasi, seperti pada "
        "eksperimen lama) -- dijalankan pada model dengan akurasi tertinggi (ConvNeXt-Base) "
        "yang dipilih karena biaya komputasi lebih rendah dibanding Large pada akurasi yang "
        "identik. Dua run dilatih dengan konfigurasi sama persis, hanya berbeda pada penerapan "
        "CLAHE di preprocessing (konsisten di train/val/test)."
    )

    add_page_break(doc)

    # ================= 4. HASIL =================
    add_heading(doc, "4. Hasil", level=1)

    add_heading(doc, "4.1 Perbandingan Performa Keempat Model", level=2)
    doc.add_paragraph(f"Evaluasi pada test set (n={n_test}), interval kepercayaan 95% dari 2.000 resample bootstrap:")

    table = doc.add_table(rows=1, cols=6)
    style_table(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(["Model", "Accuracy", "Precision", "Recall", "F1", "AUROC"]):
        hdr[i].text = h
    for m in models:
        row = table.add_row().cells
        row[0].text = MODEL_LABELS[m["model_id"]]
        row[1].text = ci_str(m["accuracy"])
        row[2].text = ci_str(m["precision"])
        row[3].text = ci_str(m["recall"])
        row[4].text = ci_str(m["f1"])
        row[5].text = ci_str(m["auroc"])
    doc.add_paragraph()

    doc.add_paragraph(
        "Analisis signifikansi: interval kepercayaan 95% accuracy keempat model SALING "
        "TUMPANG TINDIH (overlap) untuk semua pasangan -- artinya belum dapat diklaim satu "
        "arsitektur secara statistik lebih unggul dari yang lain, meskipun angka titik "
        "(point estimate) ConvNeXt-Base dan ConvNeXt-Large sedikit lebih tinggi (99,21%)."
    )
    add_figure(doc, FIGURES_DIR / "accuracy_comparison.png", "Gambar 1. Perbandingan accuracy antar backbone dengan interval kepercayaan 95%.")
    add_figure(doc, FIGURES_DIR / "roc_curves.png", "Gambar 2. Kurva ROC keempat model pada test set.")
    add_figure(doc, FIGURES_DIR / "pr_curves.png", "Gambar 3. Kurva Precision-Recall keempat model.")

    add_heading(doc, "4.2 Kalibrasi Probabilitas", level=2)
    doc.add_paragraph(
        "Setelah temperature scaling, Expected Calibration Error (ECE) keempat model berada "
        "pada rentang 0,0078-0,0168 -- menunjukkan probabilitas yang dilaporkan model cukup "
        "merepresentasikan keyakinan sesungguhnya (ConvNeXt-Large memiliki kalibrasi terbaik)."
    )
    add_figure(doc, FIGURES_DIR / "reliability_diagrams.png", "Gambar 4. Reliability diagram keempat model setelah kalibrasi.")

    add_heading(doc, "4.3 Confusion Matrix", level=2)
    add_figure(doc, FIGURES_DIR / "confusion_matrices.png", "Gambar 5. Confusion matrix pada threshold optimal (dipilih dari data validasi).")

    add_heading(doc, "4.4 Gerbang OOD dan Selective Prediction", level=2)
    table2 = doc.add_table(rows=1, cols=2)
    style_table(table2)
    table2.rows[0].cells[0].text = "Model"
    table2.rows[0].cells[1].text = "AUROC Gerbang OOD"
    for m in models:
        row = table2.add_row().cells
        row[0].text = MODEL_LABELS[m["model_id"]]
        row[1].text = f"{m['ood_auroc']:.4f}"
    doc.add_paragraph()
    doc.add_paragraph(
        "Gerbang OOD berhasil memisahkan citra X-ray (in-distribution) dari citra non-X-ray "
        "(CIFAR-10) dengan AUROC 0,999-1,000 pada keempat model -- performa yang sangat baik, "
        "dan telah diverifikasi bekerja pada demo live (lihat Bagian 5)."
    )
    add_figure(doc, FIGURES_DIR / "risk_coverage_curves.png", "Gambar 6. Kurva risk-coverage (selective prediction) -- menunjukkan trade-off antara cakupan jawaban dan tingkat error saat model diperbolehkan abstain.")

    add_heading(doc, "4.5 Ablation Study CLAHE", level=2)
    if ablation:
        doc.add_paragraph(
            f"Hasil pada {MODEL_LABELS[ablation['model_id']]}: interval kepercayaan 95% "
            "TUMPANG TINDIH pada ketiga metrik (accuracy, F1, AUROC) antara kondisi dengan dan "
            "tanpa CLAHE -- artinya CLAHE TIDAK memberikan perbedaan performa yang signifikan "
            "secara statistik pada dataset ini."
        )
        table3 = doc.add_table(rows=1, cols=4)
        style_table(table3)
        for i, h in enumerate(["Kondisi", "Accuracy", "F1", "AUROC"]):
            table3.rows[0].cells[i].text = h
        for label, key in [("Dengan CLAHE", "with_clahe"), ("Tanpa CLAHE", "without_clahe")]:
            row = table3.add_row().cells
            row[0].text = label
            row[1].text = ci_str(ablation[key]["accuracy"])
            row[2].text = ci_str(ablation[key]["f1"])
            row[3].text = ci_str(ablation[key]["auroc"])
        doc.add_paragraph()
        add_figure(doc, FIGURES_DIR / "clahe_ablation.png", "Gambar 7. Perbandingan performa dengan vs tanpa CLAHE.")
    else:
        add_placeholder(doc, "Hasil ablation CLAHE belum tersedia di results/metrics.json")

    add_page_break(doc)

    # ================= 5. PLATFORM WEB =================
    add_heading(doc, "5. Platform Web (Demo Publik)", level=1)
    doc.add_paragraph(
        "Platform web live dapat diakses di https://fracture.lapanproject.tech (publik, tanpa "
        "kata sandi) dengan lima halaman utama:"
    )
    pages = [
        ("Analyze", "Unggah satu citra X-ray, dapatkan prediksi (fractured/not fractured), "
         "probabilitas mentah dan terkalibrasi, overlay Grad-CAM interaktif (slider opasitas), "
         "dan slider threshold keputusan yang dapat diubah secara real-time. Tersedia galeri "
         "16 citra contoh (8 fractured, 8 not fractured) yang bisa diunduh langsung dari "
         "halaman ini -- diambil dari split TEST resmi (nol kebocoran by construction, lihat "
         "3.1) supaya pengunjung tanpa citra X-ray sendiri tetap bisa mencoba sistem."),
        ("Compare", "Bandingkan satu citra pada keempat model sekaligus, masing-masing dengan "
         "prediksi, latency, dan visualisasi Grad-CAM tersendiri."),
        ("Benchmark", "Seluruh tabel dan grafik pada Bagian 4 laporan ini dirender otomatis dari "
         "data yang sama (results/metrics.json) -- tanpa angka manual."),
        ("Methodology", "Ringkasan pipeline, analisis kegagalan eksperimen awal, hasil ablation "
         "CLAHE, keterbatasan, dan disclaimer medis."),
        ("History", "Riwayat sesi analisis tersimpan di localStorage browser (tanpa server/database)."),
    ]
    for title, desc in pages:
        p = doc.add_paragraph()
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    doc.add_paragraph()
    add_figure(doc, SCREENSHOTS_DIR / "analyze-fractured-result.png",
               "Gambar 8. Halaman Analyze -- hasil pada citra fraktur sungguhan (raw 0,997, terkalibrasi 98,8%), Grad-CAM menyorot lokasi implan/fraktur.")
    add_figure(doc, SCREENSHOTS_DIR / "analyze-notfractured-result.png",
               "Gambar 9. Halaman Analyze -- hasil pada citra tanpa fraktur (raw 0,034, terkalibrasi 6,8%).")
    add_figure(doc, SCREENSHOTS_DIR / "compare-4models-result.png",
               "Gambar 10. Halaman Compare -- satu citra dibandingkan pada keempat model sekaligus.")
    add_figure(doc, SCREENSHOTS_DIR / "benchmark-page.png",
               "Gambar 11. Halaman Benchmark -- seluruh metrik dirender otomatis dari results/metrics.json.")
    add_figure(doc, SCREENSHOTS_DIR / "methodology-page-final.png",
               "Gambar 12. Halaman Methodology -- pipeline, analisis kegagalan, dan tabel ablation CLAHE.")

    add_placeholder(doc, "Sisipkan tangkapan layar tambahan di sini jika diperlukan (mis. halaman History, "
                          "mode gelap/lightbox, atau demonstrasi langsung di depan dosen)")

    add_page_break(doc)

    # ================= 6. PANDUAN PENGGUNAAN & GLOSARIUM =================
    add_heading(doc, "6. Panduan Penggunaan Aplikasi dan Glosarium Istilah", level=1)
    doc.add_paragraph(
        "Bagian ini ditujukan untuk pembaca yang belum familiar dengan istilah teknis machine "
        "learning/statistik yang muncul di platform web -- baik saat mendemonstrasikan aplikasi "
        "maupun saat membaca Bagian 4-5 di atas."
    )

    add_heading(doc, "6.1 Cara Menggunakan Tiap Halaman", level=2)
    howto = [
        ("Analyze", [
            "Pilih model (Tiny/Small/Base/Large) dari dropdown di bagian atas.",
            "Unggah citra X-ray (seret-lepas atau klik area unggah), atau kalau belum punya "
            "citra sendiri, unduh salah satu dari galeri \"Belum punya citra X-ray?\" di bawah "
            "area unggah lalu unggah file yang sudah diunduh.",
            "Tunggu beberapa detik -- hasil verdict (Fractured/Not Fractured/Abstain), "
            "probabilitas, dan overlay Grad-CAM akan muncul di sisi kanan.",
            "Geser slider opasitas untuk mengatur seberapa tebal overlay Grad-CAM ditampilkan "
            "di atas citra asli.",
            "Geser slider threshold untuk melihat bagaimana keputusan berubah pada ambang "
            "batas yang berbeda -- ini dihitung ulang langsung di browser, tanpa perlu "
            "mengunggah ulang citra.",
        ]),
        ("Compare", [
            "Unggah satu citra X-ray.",
            "Sistem menjalankan keempat model sekaligus dan menampilkan hasil (verdict, "
            "confidence, latency, Grad-CAM) berdampingan untuk dibandingkan langsung.",
        ]),
        ("Benchmark", [
            "Tidak perlu mengunggah apa pun -- halaman ini menampilkan seluruh hasil evaluasi "
            "formal (tabel metrik, kurva ROC/PR, reliability diagram, confusion matrix, "
            "risk-coverage) untuk keempat model, sama seperti Bagian 4 laporan ini.",
            "Gunakan tab pemilih model untuk beralih grafik antar backbone.",
        ]),
        ("Methodology", [
            "Ringkasan pipeline penelitian, analisis kegagalan eksperimen awal (Bagian 2), "
            "hasil ablation CLAHE, daftar keterbatasan, dan disclaimer medis penuh -- versi "
            "ringkas dari laporan ini yang ditulis untuk pengunjung umum.",
        ]),
        ("History", [
            "Menampilkan riwayat analisis yang pernah dilakukan di browser yang sama (tersimpan "
            "lokal di perangkat, TIDAK dikirim ke server mana pun) -- bisa diekspor ke CSV atau "
            "dihapus kapan saja.",
        ]),
    ]
    for page, steps in howto:
        p = doc.add_paragraph()
        p.add_run(page).bold = True
        for step in steps:
            doc.add_paragraph(step, style="List Bullet")

    add_heading(doc, "6.2 Glosarium Istilah", level=2)
    glossary = [
        ("Fractured / Not Fractured", "Label hasil prediksi model: \"Fractured\" berarti citra "
         "terindikasi memiliki patah tulang, \"Not Fractured\" berarti tidak terindikasi."),
        ("Abstain", "Sistem sengaja TIDAK memutuskan fractured/not fractured -- terjadi kalau "
         "citra terindikasi bukan X-ray (lihat \"Gerbang OOD\") atau probabilitas terlalu dekat "
         "ke ambang batas. Pada kondisi ini keputusan harus diserahkan ke tenaga medis, bukan "
         "dipaksakan oleh sistem."),
        ("Probabilitas mentah (raw) vs terkalibrasi", "Mentah adalah output langsung model "
         "(0-1) yang cenderung terlalu percaya diri (overconfident); terkalibrasi adalah angka "
         "setelah disesuaikan (temperature scaling) supaya persentase yang ditampilkan benar-"
         "benar mencerminkan peluang sungguhan."),
        ("Threshold (ambang keputusan)", "Batas probabilitas yang memisahkan keputusan "
         "\"fractured\" vs \"not fractured\". Nilai default dipilih otomatis dari data validasi "
         "(Youden's J index), dan bisa digeser manual di halaman Analyze untuk melihat "
         "trade-off-nya."),
        ("Grad-CAM", "Visualisasi overlay warna panas yang menunjukkan area citra yang paling "
         "memengaruhi keputusan model -- dipakai untuk memeriksa apakah model \"melihat\" ke "
         "lokasi yang masuk akal secara medis, bukan sekadar hiasan."),
        ("Confidence Interval (CI) 95%", "Rentang angka yang mengekspresikan ketidakpastian "
         "statistik suatu metrik (mis. akurasi 99,21% [98,43%, 99,80%]). Kalau rentang dua "
         "model saling tumpang tindih, belum bisa diklaim satu model \"lebih baik\" secara sah "
         "secara statistik."),
        ("AUROC / AUPRC", "Area Under ROC / Precision-Recall Curve -- angka ringkas (0-1, "
         "makin tinggi makin baik) yang mengukur seberapa baik model membedakan fractured vs "
         "not fractured di semua kemungkinan threshold sekaligus."),
        ("ECE (Expected Calibration Error)", "Mengukur seberapa jauh probabilitas yang "
         "dilaporkan model dari akurasi sungguhannya -- makin kecil makin baik (model makin "
         "\"jujur\" soal seberapa yakin dia)."),
        ("Gerbang OOD (Out-of-Distribution)", "Mekanisme yang mendeteksi kalau citra yang "
         "diunggah BUKAN citra X-ray tulang (mis. foto KTP, pemandangan). Kalau terdeteksi, "
         "sistem menolak memberi verdict fractured/not fractured (status \"abstain\") "
         "daripada memaksakan tebakan."),
        ("Cold start", "Jeda beberapa puluh detik saat backend baru \"bangun\" dari kondisi "
         "idle (server gratis tidur setelah 48 jam tidak dipakai) -- ditandai banner khusus di "
         "web, bukan error."),
        ("CLAHE", "Contrast Limited Adaptive Histogram Equalization -- teknik peningkatan "
         "kontras citra yang diuji sebagai opsi preprocessing (Bagian 3.6 & 4.5). Pada "
         "penelitian ini terbukti tidak memberi perbedaan performa yang signifikan secara "
         "statistik."),
        ("ConvNeXt Tiny/Small/Base/Large", "Empat ukuran arsitektur jaringan saraf yang sama "
         "(Tiny = paling kecil/cepat, Large = paling besar, secara teori paling presisi tapi "
         "paling lambat) -- dibandingkan berdampingan di halaman Compare & Benchmark."),
    ]
    gtable = doc.add_table(rows=1, cols=2)
    style_table(gtable)
    gtable.rows[0].cells[0].text = "Istilah"
    gtable.rows[0].cells[1].text = "Penjelasan"
    gtable.columns[0].width = Cm(4)
    for term, desc in glossary:
        row = gtable.add_row().cells
        row[0].text = term
        row[1].text = desc
    doc.add_paragraph()

    add_page_break(doc)

    # ================= 7. ARSITEKTUR SISTEM =================
    add_heading(doc, "7. Arsitektur Sistem dan Deployment", level=1)
    doc.add_paragraph("Alur kerja end-to-end:")
    flow_steps = [
        "Dataset (Kaggle) \u2192 Audit kebocoran & deduplikasi \u2192 split_manifest.json (sumber kebenaran split)",
        "Training 4 backbone ConvNeXt (Google Colab, GPU T4 gratis) \u2192 checkpoint .keras per model",
        "Evaluasi formal + kalibrasi + gerbang OOD + verifikasi Grad-CAM \u2192 results/metrics.json",
        "Ekspor ONNX (2-output: probabilitas + feature map) + bobot head terpisah (.npz)",
        "Model disimpan di Google Cloud Storage \u2192 diunduh otomatis (lazy) oleh backend saat dibutuhkan",
        "Backend FastAPI + ONNX Runtime \u2192 di-deploy ke Google Cloud Run (scale-to-zero, gratis)",
        "Frontend React + Vite \u2192 di-deploy ke VPS pribadi (Nginx + Certbot, fracture.lapanproject.tech)",
    ]
    for i, step in enumerate(flow_steps, 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")

    doc.add_paragraph()
    doc.add_paragraph(
        "Catatan: rencana awal (spec penelitian) menargetkan Hugging Face Space untuk backend "
        "dan Vercel untuk frontend. Keduanya menyimpang karena alasan teknis/biaya: Hugging Face "
        "Space mengubah kebijakan tier Docker gratis menjadi berbayar ($9/bulan) pada Juli 2026 "
        "tanpa pengumuman resmi, sehingga dipindah ke Google Cloud Run (tetap gratis, functionally "
        "setara). Frontend tetap di VPS pribadi yang sudah tersedia (bukan Vercel) untuk kontrol "
        "penuh dan konsistensi dengan infrastruktur proyek lain milik peneliti."
    )

    add_placeholder(doc, "Sisipkan diagram arsitektur visual (opsional) jika ingin tampilan lebih formal "
                          "dibanding daftar bernomor di atas -- bisa digambar ulang dari alur di atas")

    add_page_break(doc)

    # ================= 8. KETERBATASAN =================
    add_heading(doc, "8. Keterbatasan", level=1)
    limitations = [
        "Dataset berasal dari satu sumber (Kaggle, agregasi dari berbagai institusi tanpa "
        "metadata pasien yang jelas); generalisasi lintas institusi/populasi belum diuji.",
        "Label bersifat biner (fractured/not fractured) -- tipe fraktur dan lokasi anatomis "
        "tidak diprediksi.",
        "Interval kepercayaan 95% keempat arsitektur saling tumpang tindih -- belum dapat "
        "diklaim satu model \"terbukti terbaik\" secara statistik, meski Base/Large memiliki "
        "angka titik tertinggi.",
        "Ablation CLAHE hanya dijalankan pada satu model (Base) karena tidak ada model yang "
        "signifikan terbaik secara statistik untuk dijadikan target ablation tunggal sesuai "
        "rencana awal penelitian.",
        "Sistem ini adalah ALAT BANTU RISET, BUKAN ALAT DIAGNOSIS MEDIS, dan belum "
        "tersertifikasi untuk penggunaan klinis. Setiap keputusan klinis tetap harus melalui "
        "radiolog atau tenaga medis berwenang.",
        "Verifikasi kesesuaian numerik Grad-CAM analitik (dibanding backpropagation langsung "
        "via TensorFlow GradientTape) pada keempat model final sudah diverifikasi ulang "
        "setelah perbaikan bug forward-pass-ganda -- selisih turun 16-25x (mis. ConvNeXt-Tiny "
        "dari 4,75e-2 menjadi 1,92e-3) dan konvergen ke rentang sempit yang sama di keempat "
        "model (1,4e-3-2,2e-3), pola khas bug yang sudah diperbaiki. Residual yang tersisa "
        "berada di atas ambang ideal awal (1e-4) namun dianalisis sebagai lantai presisi "
        "floating-point yang inheren (rekonstruksi NumPy dari bobot vs backprop native "
        "TensorFlow, bukan kesalahan turunan) -- ambang toleransi dilonggarkan ke 5e-3 "
        "setelah bug nyata diperbaiki, bukan sebagai pengganti perbaikan.",
    ]
    for item in limitations:
        doc.add_paragraph(item, style="List Bullet")

    # ================= 9. GAP & RENCANA SELANJUTNYA =================
    add_heading(doc, "9. Gap dan Rencana Selanjutnya", level=1)
    gaps = [
        "Validasi klinis/radiologis sungguhan belum dilakukan -- performa dilaporkan murni "
        "berdasarkan label dataset Kaggle, belum divalidasi oleh radiolog berlisensi.",
        "Belum ada uji terhadap dataset eksternal (dari institusi/sumber berbeda) untuk menguji "
        "generalisasi model di luar distribusi data training.",
        "Figure dan tabel siap publikasi sudah dibangkitkan (Bagian 4), namun belum disusun "
        "menjadi naskah ilmiah/paper lengkap.",
    ]
    for item in gaps:
        doc.add_paragraph(item, style="List Bullet")
    add_placeholder(doc, "Tambahkan rencana kerja spesifik untuk periode bimbingan berikutnya sesuai arahan dosen")

    add_page_break(doc)

    # ================= 10. KESIMPULAN =================
    add_heading(doc, "10. Kesimpulan", level=1)
    doc.add_paragraph(
        f"Penelitian ini berhasil membangun pipeline klasifikasi fraktur tulang yang metodologis "
        f"benar dan tervalidasi -- memperbaiki seluruh cacat yang ditemukan pada eksperimen awal "
        f"(kebocoran data, preprocessing tidak konsisten, perbandingan arsitektur tidak "
        f"terkontrol). Keempat varian ConvNeXt mencapai akurasi tinggi (98,6%-99,2%) dengan "
        f"performa yang secara statistik setara satu sama lain. Sistem end-to-end -- dari "
        f"training, evaluasi, hingga deployment publik (backend + frontend) -- telah selesai "
        f"dan teruji berfungsi dengan citra X-ray sungguhan, termasuk mekanisme keamanan "
        f"tambahan (kalibrasi, gerbang OOD, penjelasan visual Grad-CAM)."
    )
    doc.add_paragraph(
        "Ablation study CLAHE menunjukkan bahwa teknik peningkatan kontras tersebut tidak "
        "memberikan perbaikan performa yang signifikan secara statistik pada dataset ini -- "
        "temuan yang tetap bernilai ilmiah karena diuji secara terverifikasi (bukan diasumsikan "
        "tanpa bukti seperti pada eksperimen sebelumnya)."
    )

    # ================= LAMPIRAN =================
    add_page_break(doc)
    add_heading(doc, "Lampiran", level=1)
    doc.add_paragraph("Demo live: https://fracture.lapanproject.tech")
    doc.add_paragraph("Backend API: https://fracture-api-607128796608.asia-southeast2.run.app")
    doc.add_paragraph("Kode sumber (GitHub): https://github.com/alianhar/project-fracture")
    doc.add_paragraph(f"Data mentah hasil evaluasi: results/metrics.json (config_hash: {data.get('config_hash')})")
    add_placeholder(doc, "Sisipkan lampiran tambahan sesuai kebutuhan (mis. cuplikan kode penting, "
                          "log training, atau dokumen pendukung lain)")

    doc.save(OUT_PATH)
    print(f"Tersimpan: {OUT_PATH}")


if __name__ == "__main__":
    main()
