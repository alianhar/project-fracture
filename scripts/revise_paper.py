"""
Revisi docs/LAPORAN-PENELITIAN-Bone-Fracture.docx supaya SINKRON dengan
pipeline yang sudah diaudit & diperbaiki (results/metrics.json), dan
membersihkan kontaminasi topik "penyakit paru-paru / 5 kelas" yang
tertinggal dari template/sumber lain.

Latar belakang lengkap ada di riwayat percakapan sesi ini (2026-08-23):
paper versi lama melaporkan hasil eksperimen SEBELUM audit dataset --
3 dari 4 model (Large/Base/Tiny) collapse ke ~48-53% akurasi (bug
preprocessing train!=test), dan 1 model (Small) kebetulan 98,6% tapi
lewat data yang bocor 34,84% (duplikat train<->test) -- BUKAN karena
model generalisasi. Kedua bug ini sudah ditemukan & diperbaiki di
pipeline saat ini (lihat CLAUDE.md bagian [2] dan [3]).

TIDAK menyentuh: cover, lembar pengesahan, surat pernyataan plagiasi,
abstrak, kata pengantar, BAB I, BAB II (tinjauan pustaka), daftar
pustaka -- hanya BAB III (metode) + BAB IV (hasil) + BAB V (kesimpulan)
yang direvisi, karena itu bagian yang datanya tidak sinkron.

Section kalkulasi manual konvolusi/ReLU/perceptron (BAB IV S4.4-4.5,
~150 paragraf) SENGAJA TIDAK DIHAPUS per keputusan eksplisit user --
hanya diberi catatan metodologis (disclaimer) bahwa itu ilustrasi
generik, bukan transkrip pelatihan sungguhan.

Jalankan: python scripts/revise_paper.py
Input : docs/LAPORAN-PENELITIAN-Bone-Fracture.docx (ditimpa di tempat)
Backup: docs/LAPORAN-PENELITIAN-Bone-Fracture.backup-<timestamp>.docx
        dibuat otomatis SEBELUM ditimpa.
"""

import json
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt

REPO_ROOT = Path(__file__).resolve().parent.parent
DOCX_PATH = REPO_ROOT / "docs" / "LAPORAN-PENELITIAN-Bone-Fracture.docx"
METRICS_PATH = REPO_ROOT / "results" / "metrics.json"
PAPER_FIG_DIR = REPO_ROOT / "results" / "figures" / "paper"
RUNS_DIR = REPO_ROOT / "fracture-runs"

MODEL_ORDER = ["tiny", "small", "base", "large"]
MODEL_LABELS = {"tiny": "ConvNeXt-Tiny", "small": "ConvNeXt-Small", "base": "ConvNeXt-Base", "large": "ConvNeXt-Large"}
MODEL_RUN_DIRS = {
    "tiny": "tiny_a817fd5a",
    "small": "small_4fdac66d",
    "base": "base_cf7600d6",
    "large": "large_f6062dcc",
}
MODEL_PHASE1_EPOCHS = {"tiny": 30, "small": 30, "base": 23, "large": 30}

# Angka training terverifikasi manual dari history_phase2.csv (baris dgn
# val_loss minimum -- kriteria ModelCheckpoint sungguhan, monitor="val_loss"),
# lihat riwayat sesi. Ditulis literal di sini (bukan dihitung ulang saat
# run) supaya skrip ini tidak butuh TensorFlow/pandas -- angka sudah
# diverifikasi silang dgn CLAUDE.md "Hasil akhir (total epoch ASLI, val
# terbaik)".
TRAIN_METRICS = {
    "tiny":  {"total_epochs": 51, "train_acc": 0.9979, "train_loss": 0.0079, "val_acc": 0.9861, "val_loss": 0.0606},
    "small": {"total_epochs": 58, "train_acc": 1.0000, "train_loss": 0.0029, "val_acc": 0.9861, "val_loss": 0.0391},
    "base":  {"total_epochs": 41, "train_acc": 0.9992, "train_loss": 0.0033, "val_acc": 0.9921, "val_loss": 0.0476},
    "large": {"total_epochs": 40, "train_acc": 0.9919, "train_loss": 0.0215, "val_acc": 0.9881, "val_loss": 0.0408},
}

# Prediksi sungguhan dari backend live (Cloud Run), diambil thd
# web/public/samples/{fractured,not_fractured}/sample-01.jpg, 2026-08-23.
SAMPLE_PREDICTIONS = {
    "tiny":  {"fractured_conf": 0.9962, "notfrac_conf": 1 - 0.012417},
    "small": {"fractured_conf": 0.999975, "notfrac_conf": 1 - 0.195207},
    "base":  {"fractured_conf": 0.999862, "notfrac_conf": 1 - 0.025153},
    "large": {"fractured_conf": 0.999880, "notfrac_conf": 1 - 0.000123},
}


def backup():
    ts = datetime.now().strftime("%Y%m%d-%H%M%S")
    backup_path = DOCX_PATH.with_name(f"LAPORAN-PENELITIAN-Bone-Fracture.backup-{ts}.docx")
    shutil.copy2(DOCX_PATH, backup_path)
    print(f"Backup: {backup_path.name}")
    return backup_path


def set_text(doc, idx: int, new_text: str) -> None:
    """Ganti isi paragraf tanpa merusak style paragraf (Heading/Normal/dst).
    Menyalin format run pertama (font/size/bold/italic) kalau ada, supaya
    tidak mendadak jadi default Times New Roman polos di tengah paragraf
    yang tadinya calibri/dst."""
    p = doc.paragraphs[idx]
    base_font = None
    if p.runs:
        base_font = p.runs[0].font
    for r in list(p.runs):
        r.text = ""
    if p.runs:
        p.runs[0].text = new_text
    else:
        run = p.add_run(new_text)
        if base_font is not None:
            run.font.name = base_font.name
            run.font.size = base_font.size
            run.font.bold = base_font.bold
            run.font.italic = base_font.italic


def set_cell(table, row: int, col: int, text: str, bold: bool = False) -> None:
    cell = table.cell(row, col)
    cell.text = text
    if bold:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True


def replace_image_paragraph(doc, idx: int, image_path: Path, width_cm: float) -> None:
    """Paragraf idx berisi HANYA gambar (dikonfirmasi manual saat audit
    struktur dokumen) -- hapus run/drawing lama, tambahkan gambar baru."""
    p = doc.paragraphs[idx]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))


def insert_paragraphs_before(anchor_para, blocks: list[tuple[str, str | None]]):
    """blocks: list of (text, style_name_or_None). Insert sebelum anchor_para,
    urutan sesuai list (paling atas duluan)."""
    for text, style in blocks:
        new_p = anchor_para.insert_paragraph_before(text, style=style)
        new_p.paragraph_format.space_after = Pt(8)


def main():
    with open(METRICS_PATH, encoding="utf-8") as f:
        metrics = json.load(f)
    models_by_id = {m["model_id"]: m for m in metrics["models"]}

    backup()
    doc = Document(DOCX_PATH)

    # ================================================================
    # PASS 1 -- bersihkan kontaminasi "penyakit paru-paru / 5 kelas"
    # (teks murni, indeks paragraf stabil karena belum ada insert/delete)
    # ================================================================
    set_text(doc, 338,
        "Meskipun telah banyak model deep learning yang digunakan untuk deteksi "
        "patah tulang, seperti ResNet, VGG, dan DenseNet, sebagian besar masih "
        "menghadapi kendala dalam hal efisiensi pemrosesan, akurasi klasifikasi, "
        "serta kemampuan generalisasi pada dataset yang kompleks dan tidak "
        "seimbang. Dengan hadirnya arsitektur ConvNeXt, yang menggabungkan "
        "kekuatan CNN klasik dan penyempurnaan dari pendekatan transformer, "
        "terdapat potensi untuk mengatasi keterbatasan tersebut. Namun, "
        "implementasi ConvNeXt dalam domain deteksi patah tulang melalui citra "
        "X-ray masih belum banyak dieksplorasi, sehingga efektivitas dan "
        "efisiensinya perlu diteliti lebih lanjut.")

    set_text(doc, 347,
        "Analisis data adalah proses penting dalam penelitian klasifikasi patah "
        "tulang yang menggunakan Convolutional Neural Network (CNN). Proses ini "
        "bertujuan untuk memahami karakteristik data, mempersiapkan data untuk "
        "pelatihan model, serta meningkatkan akurasi dan efisiensi model yang "
        "dihasilkan. Analisis data membantu mengidentifikasi pola, membersihkan "
        "data dari potensi anomali, dan memastikan data siap digunakan dalam "
        "proses pengolahan dengan CNN. Berikut adalah langkah-langkah analisis "
        "data yang dapat dilakukan:")

    set_text(doc, 348,
        "Mengidentifikasi sumber dataset, yaitu citra X-ray tulang pasien "
        "dengan dua kategori kondisi (patah/fractured dan tidak patah/not "
        "fractured), serta memahami jenis data yang dimiliki, misalnya format "
        "gambar, jumlah kelas, serta label yang tersedia.")

    set_text(doc, 349,
        "Menghitung jumlah total citra per kelas untuk memastikan distribusi "
        "yang seimbang dan melakukan visualisasi sampel citra untuk memahami "
        "variasi antar kelas (fractured/not fractured) serta kualitas gambar.")

    set_text(doc, 351,
        "Memastikan dataset memiliki distribusi yang seimbang antar kelas "
        "fractured/not fractured agar model tidak bias terhadap kelas tertentu.")

    set_text(doc, 353,
        "Jika diperlukan, mengekstrak fitur tambahan dari citra X-ray untuk "
        "meningkatkan akurasi model CNN.")

    set_text(doc, 359,
        "Flowchart pada Gambar (3.2) tersebut menggambarkan alur proses "
        "pelatihan model klasifikasi citra menggunakan arsitektur ConvNeXt. "
        "Dimulai dari tahap input data, dilakukan pre-processing seperti "
        "normalisasi, augmentasi, dan perubahan ukuran (resize). Data kemudian "
        "dimasukkan ke dalam model ConvNeXt yang telah dilatih sebelumnya "
        "(pre-trained), dimulai dengan layer konvolusi awal dan beberapa blok "
        "ConvNeXt yang diselingi oleh downsampling kemudian diikuti oleh proses "
        "global average pooling, layer normalization, linear, dan sigmoid yang "
        "menghasilkan output klasifikasi BINER (fractured / not fractured) -- "
        "CATATAN: label \"5 kelas\"/\"softmax\" pada diagram Gambar 3.2 adalah "
        "sisa template generik dan TIDAK merepresentasikan output sungguhan "
        "penelitian ini, yang biner (1 neuron output, aktivasi sigmoid). "
        "Setelah itu, model dikompilasi dan dilatih, lalu hasil pelatihan "
        "dievaluasi untuk melihat apakah akurasi tinggi dan tidak terjadi "
        "overfitting. Jika ya, performa dievaluasi lebih lanjut dan metrik "
        "akhir ditampilkan; jika tidak, proses pelatihan diulang. Flowchart "
        "ini berakhir pada tahap evaluasi performa dan hasil akhir metrik "
        "evaluasi.")

    set_text(doc, 362,
        "Proses implementasi model dilakukan melalui platform Google Colab, "
        "yang memungkinkan eksekusi kode Python secara online tanpa perlu "
        "instalasi di komputer lokal, serta memanfaatkan GPU yang disediakan "
        "oleh Google Colab untuk mempercepat proses pelatihan model. Bahasa "
        "pemrograman Python dipilih karena kompatibel dengan berbagai pustaka "
        "deep learning, termasuk TensorFlow. Dalam kajian ini, TensorFlow "
        "digunakan sebagai kerangka kerja utama dalam pengembangan dan "
        "pelatihan model ConvNeXt untuk klasifikasi citra X-ray tulang.")

    set_text(doc, 365,
        "Pengujian dalam deteksi patah tulang menggunakan arsitektur ConvNeXt "
        "merupakan tahap evaluasi yang bertujuan untuk mengukur kinerja model "
        "yang telah dilatih terhadap data yang belum pernah dilihat "
        "sebelumnya. Analisis confusion matrix dilakukan untuk memahami "
        "distribusi prediksi di setiap kelas dan mengidentifikasi potensi "
        "kesalahan klasifikasi, seperti prediksi salah terhadap kondisi "
        "fractured atau not fractured.")

    set_text(doc, 366,
        "Pengujian ini juga berperan penting dalam mengevaluasi kemampuan "
        "generalisasi model terhadap data baru, serta memverifikasi apakah "
        "model mengalami overfitting atau tidak. Hasil dari tahap pengujian "
        "memberikan gambaran menyeluruh tentang keandalan dan efektivitas "
        "model dalam mendeteksi patah tulang, serta menentukan sejauh mana "
        "model dapat diandalkan untuk diaplikasikan dalam dunia medis, "
        "seperti dalam sistem pendukung diagnosis berbasis citra X-Ray.")

    set_text(doc, 383,
        "Pada bagian ini dibahas hasil prediksi model ConvNeXt terhadap data "
        "uji berupa citra X-ray tulang. Pembahasan mencakup kelas hasil "
        "prediksi dan nilai probabilitas (confidence score) dari masing-"
        "masing kelas.")

    set_text(doc, 397,
        "Analisis data merupakan tahapan yang sangat penting dalam penelitian "
        "ini. Pada tahap ini, peneliti mengumpulkan data citra X-ray tulang "
        "yang akan berfungsi sebagai input data. Prosedur pengumpulan "
        "dilakukan melalui platform Kaggle dengan mengunjungi URL "
        "https://www.kaggle.com/datasets/usman44m/bone-fracture-x-ray-dataset. "
        "Dari sumber tersebut, klaim awal Kaggle adalah 10.587 gambar -- "
        "TETAPI audit deduplikasi (dijelaskan pada subbab 4.1.1) menemukan "
        "34,84% di antaranya adalah duplikat byte-identik yang tersebar "
        "sembarangan ke folder train/val/test, sehingga jumlah CITRA UNIK "
        "sungguhan hanya 3.370. Dataset yang digunakan dalam penelitian ini "
        "terbagi menjadi dua kelas, yakni Fractured dan Not Fractured. "
        "Gambar-gambar yang diambil memiliki dimensi ukuran yang bervariasi "
        "dan berformat png dan jpg. Jumlah data latih, data uji, dan data "
        "validasi (dari 3.370 citra unik, BUKAN dari 10.587 klaim awal) dapat "
        "dilihat dalam Tabel 4.1.")

    set_text(doc, 402, "Normalisasi Citra")
    set_text(doc, 403,
        "ConvNeXt (implementasi Keras Applications yang dipakai penelitian "
        "ini) menerima citra RGB mentah pada rentang [0, 255] TANPA rescaling "
        "manual -- fungsi preprocess_input bawaan ConvNeXt dikonfirmasi "
        "bersifat identity (tidak melakukan transformasi apa pun secara "
        "numerik), karena normalisasi sudah menjadi bagian dari layer di "
        "dalam arsitektur model itu sendiri. Implikasi praktis: TIDAK ADA "
        "rescaling piksel manual (mis. ke [-1, 1] atau [0, 1]) yang "
        "diterapkan di luar model pada penelitian ini -- deviasi dari "
        "praktik CNN klasik yang umumnya membutuhkan normalisasi eksplisit.")

    set_text(doc, 405,
        "CLAHE digunakan untuk meningkatkan kontras lokal pada gambar, "
        "terutama di area dengan pencahayaan rendah atau kontras yang tidak "
        "merata. Teknik ini membagi gambar menjadi beberapa wilayah kecil "
        "(tile) dan melakukan equalisasi histogram pada setiap wilayah, "
        "dengan batas kliping untuk mencegah peningkatan noise yang "
        "berlebihan. CATATAN PENTING: CLAHE TIDAK diterapkan sebagai bagian "
        "dari pipeline preprocessing utama keempat model (Tiny/Small/Base/"
        "Large) pada penelitian ini -- CLAHE hanya diuji secara terpisah "
        "sebagai ablation study pada ConvNeXt-Base (lihat subbab 4.5), dan "
        "hasilnya menunjukkan TIDAK ADA perbedaan performa yang signifikan "
        "secara statistik (interval kepercayaan 95% tumpang tindih).")

    set_text(doc, 413,
        "Rotasi diterapkan dengan memutar citra terhadap titik pusatnya, "
        "baik ke arah searah maupun berlawanan jarum jam. Pada dataset, "
        "gambar diputar secara acak dengan batas sudut maksimum ±15 "
        "derajat. Teknik ini bertujuan agar model tetap mampu mengenali "
        "objek meskipun citra diambil dari sudut pandang yang berbeda.")

    set_text(doc, 415,
        "Teknik zoom diterapkan dengan memperbesar citra secara acak hingga "
        "15% dari ukuran aslinya. Pendekatan ini memungkinkan model "
        "mempelajari fitur objek dari skala yang lebih dekat serta "
        "meningkatkan ketahanan model terhadap variasi ukuran atau perubahan "
        "skala pada gambar. Augmentasi (flip/rotasi/zoom) HANYA diterapkan "
        "pada data training -- data validasi dan uji SELALU tanpa augmentasi, "
        "supaya evaluasi mencerminkan citra asli, bukan versi yang "
        "dimodifikasi.")

    set_text(doc, 420,
        "Setelah tahap pre-processing citra selesai, selanjutnya akan "
        "dilakukan pemodelan arsitektur ConvNeXt. Tahapan ini bertujuan "
        "untuk membangun model deep learning berbasis arsitektur ConvNeXt "
        "yang mampu melakukan klasifikasi BINER (fractured / not fractured) "
        "berdasarkan citra X-Ray. Pemodelan dilakukan dengan mempertimbangkan "
        "efisiensi, kemampuan generalisasi, dan interpretabilitas model.")

    set_text(doc, 754,
        "Proses pelatihan model dilakukan dalam DUA FASE (bukan satu fase "
        "tunggal): Fase 1 -- backbone ConvNeXt dibekukan (frozen), hanya "
        "head klasifikasi yang dilatih, maksimum 30 epoch, learning rate "
        "1e-4. Fase 2 -- sebagian lapisan terakhir backbone dibuka "
        "(fine-tuning) dengan learning rate lebih kecil (1e-5), memakai "
        "EarlyStopping (monitor val_loss, patience 8 epoch, "
        "restore_best_weights) supaya training berhenti otomatis begitu "
        "model berhenti membaik -- bukan jumlah epoch tetap yang "
        "ditentukan di muka. Optimizer yang dipilih adalah Adam, yang "
        "dikenal efektif dalam menangani dataset yang besar dan kompleks, "
        "serta dapat beradaptasi secara dinamis dengan learning rate yang "
        "berbeda. Adam bekerja dengan menggabungkan keunggulan dari dua "
        "metode sebelumnya, yaitu RMSprop dan Stochastic Gradient Descent "
        "(SGD) dengan momentum (Ramadhan & Hernadi, 2025).")

    set_text(doc, 760,
        "Hasil kinerja model ConvNeXt dalam mengklasifikasikan patah tulang "
        "berdasarkan citra X-ray dievaluasi menggunakan berbagai metrik, "
        "seperti akurasi, presisi, recall, dan F1-score. Evaluasi ini "
        "dilakukan setelah proses pelatihan selesai, dengan menggunakan "
        "data testing yang tidak pernah dilihat oleh model sebelumnya. "
        "Tujuannya adalah untuk mengetahui seberapa baik model dapat "
        "menggeneralisasi terhadap data baru. Rumus untuk masing-masing "
        "metrik evaluasi yang diterapkan dalam studi ini disajikan pada "
        "rumus berikut.")

    print("Pass 1 selesai: pembersihan kontaminasi paru-paru/5-kelas.")

    # ================================================================
    # PASS 2 -- Tabel 4.1 Jumlah Dataset (real audit numbers)
    # ================================================================
    t = doc.tables[2]
    while len(t.rows) > 1:
        t._tbl.remove(t.rows[-1]._tr)
    for i in range(len(t.columns)):
        t.cell(0, i).text = ["Kelas", "Data Training", "Data Validation", "Data Testing"][i]
    rows = [
        ("Fractured", "983 citra", "210 citra", "212 citra"),
        ("Not Fractured", "1.375 citra", "294 citra", "296 citra"),
        ("TOTAL (unik, setelah dedup)", "2.358 citra", "504 citra", "508 citra"),
    ]
    for r in rows:
        row = t.add_row().cells
        for i, v in enumerate(r):
            row[i].text = v
    for c in t.rows[0].cells:
        for p in c.paragraphs:
            for run in p.runs:
                run.bold = True
    print("Pass 2 selesai: Tabel 4.1 Jumlah Dataset.")

    # ================================================================
    # PASS 3 -- Tabel Hyperparameter (indeks 3)
    # ================================================================
    t = doc.tables[3]
    hp_rows = [
        ("Function", "Value"),
        ("Optimizer", "Adam (2 fase: lr 1e-4 lalu 1e-5)"),
        ("Learning rate", "Fase 1: 0.0001 | Fase 2: 0.00001"),
        ("Loss", "binary_crossentropy"),
        ("Metric", "Accuracy"),
        ("Epochs", "Variabel per model (EarlyStopping, val_loss, patience=8) -- lihat Tabel 4.2"),
    ]
    for i, (k, v) in enumerate(hp_rows):
        set_cell(t, i, 0, k, bold=(i == 0))
        set_cell(t, i, 1, v, bold=(i == 0))
    print("Pass 3 selesai: Tabel Hyperparameter.")

    # ================================================================
    # PASS 4 -- teks interpretasi training curve per model (real numbers)
    # ================================================================
    training_para_idx = {"large": 776, "base": 781, "tiny": 786, "small": 791}
    for mid, idx in training_para_idx.items():
        tm = TRAIN_METRICS[mid]
        set_text(doc, idx,
            f"Berdasarkan visualisasi ini (Gambar training curve sungguhan, "
            f"bukan ilustrasi), model {MODEL_LABELS[mid]} dilatih selama "
            f"{tm['total_epochs']} epoch total (kedua fase, berhenti otomatis "
            f"via EarlyStopping). Akurasi training akhir mencapai "
            f"{tm['train_acc']*100:.2f}% dan akurasi validasi "
            f"{tm['val_acc']*100:.2f}% -- keduanya berdekatan, mengindikasikan "
            f"TIDAK ada overfitting parah. Loss training akhir "
            f"{tm['train_loss']:.4f} dan loss validasi {tm['val_loss']:.4f}. "
            f"Model checkpoint terbaik (best.keras) dipilih otomatis pada "
            f"epoch dengan val_loss terendah, bukan epoch terakhir.")
    print("Pass 4 selesai: teks training curve per model.")

    # ================================================================
    # PASS 5 -- Tabel Precision/Recall/F1/Akurasi per model (indeks 4-7)
    # ================================================================
    table_idx_by_model = {"large": 4, "base": 5, "tiny": 6, "small": 7}
    for mid, tidx in table_idx_by_model.items():
        m = models_by_id[mid]
        cm = m["confusion_matrix"]
        tp, fp, tn, fn = cm["tp"], cm["fp"], cm["tn"], cm["fn"]
        prec_f, rec_f = tp / (tp + fp), tp / (tp + fn)
        f1_f = 2 * prec_f * rec_f / (prec_f + rec_f)
        prec_nf, rec_nf = tn / (tn + fn), tn / (tn + fp)
        f1_nf = 2 * prec_nf * rec_nf / (prec_nf + rec_nf)
        acc = (tp + tn) / (tp + tn + fp + fn)
        t = doc.tables[tidx]
        set_cell(t, 1, 2, f"{prec_f*100:.2f}%")
        set_cell(t, 1, 3, f"{rec_f*100:.2f}%")
        set_cell(t, 1, 4, f"{f1_f*100:.2f}%")
        set_cell(t, 1, 5, f"{acc*100:.2f}%")
        set_cell(t, 2, 2, f"{prec_nf*100:.2f}%")
        set_cell(t, 2, 3, f"{rec_nf*100:.2f}%")
        set_cell(t, 2, 4, f"{f1_nf*100:.2f}%")
        set_cell(t, 2, 5, f"{acc*100:.2f}%")
    print("Pass 5 selesai: 4 tabel Precision/Recall/F1/Akurasi per model.")

    # ================================================================
    # PASS 6 -- teks interpretasi ROC per model (real AUROC)
    # ================================================================
    roc_para_idx = {"large": 834, "base": 877, "tiny": 918, "small": 960}
    for mid, idx in roc_para_idx.items():
        auroc = models_by_id[mid]["auroc"]["point"]
        set_text(doc, idx,
            f"Analisis kurva ROC pada Gambar Kurva ROC {MODEL_LABELS[mid]} "
            f"menunjukkan kemampuan diskriminatif yang SANGAT TINGGI, dengan "
            f"AUROC {auroc:.4f} (dari maksimum 1,0) pada test set (n=508, "
            f"nol kebocoran train/test by construction). Kurva mendekati "
            f"pojok kiri-atas, jauh di atas garis baseline acak (AUROC=0,5) -- "
            f"model mampu membedakan citra fractured dan not fractured hampir "
            f"sempurna pada data uji yang belum pernah dilihat model.")
    print("Pass 6 selesai: teks interpretasi ROC per model.")

    # ================================================================
    # PASS 7 -- gambar: training curve, confusion matrix, ROC (real)
    # ================================================================
    img_idx_by_model_training = {"large": 774, "base": 779, "tiny": 784, "small": 789}
    for mid, idx in img_idx_by_model_training.items():
        img = RUNS_DIR / MODEL_RUN_DIRS[mid] / "training_curve.png"
        if img.exists():
            replace_image_paragraph(doc, idx, img, width_cm=13)
        else:
            print(f"  [!] training_curve.png tidak ditemukan utk {mid}: {img}")

    img_idx_by_model_cm = {"large": 795, "base": 838, "tiny": 879, "small": 921}
    for mid, idx in img_idx_by_model_cm.items():
        img = PAPER_FIG_DIR / f"confusion_{mid}.png"
        replace_image_paragraph(doc, idx, img, width_cm=9)

    img_idx_by_model_roc = {"large": 832, "base": 875, "tiny": 916, "small": 958}
    for mid, idx in img_idx_by_model_roc.items():
        img = PAPER_FIG_DIR / f"roc_{mid}.png"
        replace_image_paragraph(doc, idx, img, width_cm=9)

    print("Pass 7 selesai: 12 gambar diganti (4 training curve + 4 confusion matrix + 4 ROC).")

    # ================================================================
    # PASS 8 -- Tabel 4.9 Hasil Prediksi (indeks 8) -- confidence real
    # ================================================================
    t = doc.tables[8]
    row_data = [
        ("1", "Fractured", "ConvNeXt Tiny", f"{SAMPLE_PREDICTIONS['tiny']['fractured_conf']*100:.2f}%"),
        ("2", "Not Fractured", "ConvNeXt Tiny", f"{SAMPLE_PREDICTIONS['tiny']['notfrac_conf']*100:.2f}%"),
        ("3", "Fractured", "ConvNeXt Small", f"{SAMPLE_PREDICTIONS['small']['fractured_conf']*100:.2f}%"),
        ("4", "Not Fractured", "ConvNeXt Small", f"{SAMPLE_PREDICTIONS['small']['notfrac_conf']*100:.2f}%"),
        ("5", "Fractured", "ConvNeXt Base", f"{SAMPLE_PREDICTIONS['base']['fractured_conf']*100:.2f}%"),
        ("6", "Not Fractured", "ConvNeXt Base", f"{SAMPLE_PREDICTIONS['base']['notfrac_conf']*100:.2f}%"),
        ("7", "Fractured", "ConvNeXt Large", f"{SAMPLE_PREDICTIONS['large']['fractured_conf']*100:.2f}%"),
        ("8", "Not Fractured", "ConvNeXt Large", f"{SAMPLE_PREDICTIONS['large']['notfrac_conf']*100:.2f}%"),
    ]
    for r, (no, kelas, model, conf) in enumerate(row_data, start=1):
        set_cell(t, r, 0, no)
        set_cell(t, r, 2, kelas)
        set_cell(t, r, 3, model)
        set_cell(t, r, 4, conf)
    print("Pass 8 selesai: Tabel Hasil Prediksi (confidence dari backend live).")

    # ================================================================
    # PASS 9 -- BAB V Kesimpulan
    # ================================================================
    set_text(doc, 972,
        "Berdasarkan hasil penelitian mengenai analisis kinerja klasifikasi "
        "patah tulang pada citra X-ray menggunakan arsitektur ConvNeXt, dapat "
        "disimpulkan bahwa penerapan metode deep learning berbasis "
        "Convolutional Neural Network mampu memberikan performa yang sangat "
        "baik dalam mendeteksi patah tulang. Klaim awal Kaggle menyebut "
        "10.587 citra, TETAPI audit deduplikasi menemukan hanya 3.370 citra "
        "yang benar-benar unik (34,84% sisanya adalah duplikat byte-identik "
        "yang tersebar sembarangan ke train/val/test) -- dataset final "
        "penelitian ini memakai 3.370 citra unik tersebut, dibagi ulang "
        "70/15/15 (train/val/test) pada level klaster duplikat supaya "
        "kebocoran data menjadi NOL secara konstruksi, bukan sekadar di "
        "bawah ambang toleransi.")

    set_text(doc, 973,
        "Tahapan preprocessing (resize 224x224, augmentasi flip/rotasi/zoom "
        "pada data training) berhasil mempersiapkan data untuk pelatihan "
        "yang efektif. CLAHE diuji SECARA TERPISAH sebagai ablation study "
        "pada ConvNeXt-Base dan terbukti TIDAK memberikan perbedaan performa "
        "yang signifikan secara statistik (interval kepercayaan 95% saling "
        "tumpang tindih pada accuracy, F1, dan AUROC) -- sehingga CLAHE "
        "TIDAK diterapkan sebagai bagian dari pipeline preprocessing utama "
        "keempat model.")

    print("Pass 9 selesai: BAB V Kesimpulan.")

    # ================================================================
    # PASS 10 (INSERTIONS -- dilakukan TERAKHIR, pakai referensi objek
    # paragraf, bukan indeks, supaya aman dari pergeseran indeks)
    # ================================================================

    # 10a. Disclaimer sebelum "4.4 Tahap Pemodelan ConvNeXt" (section
    # kalkulasi manual konvolusi/ReLU/perceptron, S4.4-4.5, SENGAJA
    # dibiarkan apa adanya per keputusan eksplisit user -- diberi catatan
    # metodologis di sini, bukan dihapus).
    anchor_44 = doc.paragraphs[419]  # heading "4.4 Tahap Pemodelan ConvNeXt"
    insert_paragraphs_before(anchor_44, [
        ("CATATAN METODOLOGIS -- BACA SEBELUM LANJUT KE SUBBAB 4.4-4.5:", "Heading 3"),
        (
            "Uraian teknis pada subbab 4.4 dan 4.5 berikut ini (perhitungan "
            "konvolusi piksel-demi-piksel per channel RGB, penerapan ReLU "
            "sel-demi-sel, average pooling, hingga contoh perceptron dengan "
            "bobot w/v numerik) adalah ILUSTRASI PEDAGOGIS GENERIK tentang "
            "cara kerja lapisan CNN secara konseptual -- nilai piksel dan "
            "bobot pada contoh tersebut bersifat HIPOTETIS/CONTOH, BUKAN "
            "diambil dari data citra atau bobot model ConvNeXt sungguhan "
            "pada penelitian ini.",
            "Normal",
        ),
        (
            "Pelatihan ConvNeXt yang SUNGGUHAN pada penelitian ini TIDAK "
            "dihitung manual seperti pada contoh berikut. Pelatihan "
            "dilakukan sepenuhnya lewat differensiasi otomatis (automatic "
            "differentiation / backpropagation) melalui framework "
            "TensorFlow/Keras, yang menghitung gradien terhadap JUTAAN "
            "parameter model sekaligus dalam satu langkah komputasi -- "
            "bukan lewat kalkulasi manual satu-per-satu seperti pada "
            "contoh perceptron sederhana di bawah, yang hanya cocok untuk "
            "menjelaskan KONSEP dasar backpropagation pada jaringan "
            "berukuran sangat kecil (beberapa neuron), bukan arsitektur "
            "ConvNeXt yang memiliki puluhan juta parameter.",
            "Normal",
        ),
        (
            "Hasil pelatihan dan evaluasi model yang SUNGGUHAN -- grafik "
            "akurasi/loss training real per model, confusion matrix, dan "
            "kurva ROC dari test set sungguhan -- disajikan pada subbab 4.6 "
            "\"Evaluasi Performa\" dan seterusnya, dibangkitkan langsung "
            "dari results/metrics.json (bukan angka manual/contoh).",
            "Normal",
        ),
    ])

    # 10b. Section baru "Audit Dataset & Alasan Meninggalkan Eksperimen
    # Awal" -- disisipkan sebelum heading "4.2 Preprocessing Data" (skrg
    # sudah bergeser krn insersi di atas TIDAK memengaruhi ini -- objek
    # paragraf diambil dari referensi asli sebelum insersi manapun, aman).
    anchor_42 = doc.paragraphs[400]  # "4.2 Preprocessing Data"
    insert_paragraphs_before(anchor_42, [
        ("4.1.1 Audit Dataset dan Alasan Revisi Metodologi", "Heading 2"),
        (
            "Sebelum pipeline final ini dibangun, dilakukan audit menyeluruh "
            "terhadap eksperimen awal (percobaan pertama dengan dataset "
            "\"apa adanya\" dari folder train/val/test Kaggle) yang menemukan "
            "dua cacat metodologis fatal. Bagian ini menjelaskan temuan "
            "tersebut secara rinci, karena angka yang dilaporkan pada bagian "
            "4.6-4.9 laporan ini BERBEDA SIGNIFIKAN dari eksperimen awal, "
            "dan perbedaan itu perlu dipahami alasannya, bukan dianggap "
            "kontradiksi yang tidak dijelaskan.",
            "Normal",
        ),
        ("Temuan 1: Kebocoran Data (Data Leakage) 34,84%", "Heading 3"),
        (
            "Audit hash konten (MD5) terhadap seluruh 10.587 file yang "
            "diklaim Kaggle menemukan bahwa 96,3% dari 508 citra pada folder "
            "test memiliki duplikat byte-identik (hash MD5 sama persis) di "
            "folder train. Dari klaim 10.587 file, hanya 3.370 citra yang "
            "benar-benar unik -- sisanya adalah salinan yang tersebar "
            "sembarangan oleh uploader dataset ke folder train/val/test "
            "TANPA deduplikasi.",
            "Normal",
        ),
        (
            "Analogi: bayangkan seorang guru memberi 100 soal latihan kepada "
            "murid, lalu saat ujian, 96 dari 100 soal ujian TERNYATA SAMA "
            "PERSIS dengan soal latihan yang sudah dihafal jawabannya oleh "
            "murid tersebut. Nilai ujian yang tinggi dalam situasi ini BUKAN "
            "bukti murid memahami materi -- itu bukti murid menghafal "
            "jawaban. Persis seperti itulah yang terjadi pada model yang "
            "dievaluasi memakai folder test Kaggle asli: model bisa saja "
            "\"mengingat\" gambar yang sama persis yang pernah dilihatnya "
            "saat training, bukan benar-benar belajar mengenali pola fraktur "
            "pada citra X-ray yang belum pernah dilihat.",
            "Normal",
        ),
        ("Temuan 2: Preprocessing Data Latih Tidak Konsisten dengan Data Uji", "Heading 3"),
        (
            "Pada eksperimen awal, generator data latih/validasi memakai "
            "rescale=1/255 (piksel diskalakan ke rentang [0, 1]), sementara "
            "generator data uji memakai fungsi preprocess_input bawaan "
            "ConvNeXt yang TERNYATA tidak melakukan transformasi numerik "
            "apa pun (identity/no-op) -- normalisasi ConvNeXt sudah menjadi "
            "bagian dari layer di dalam model itu sendiri, BUKAN langkah "
            "terpisah sebelum data masuk model. Akibatnya, tiga dari empat "
            "model diuji pada skala nilai piksel yang BERBEDA dari skala "
            "yang dipakai saat model itu dilatih.",
            "Normal",
        ),
        (
            "Analogi: bayangkan seseorang diajari membaca peta yang "
            "berskala 1:100.000 (1 cm di peta = 1 km jarak sungguhan), "
            "lalu tanpa diberi tahu, dia diuji memakai peta berskala "
            "1:25.000 (1 cm = 250 m). Orang tersebut akan tersesat dan "
            "salah menghitung jarak -- BUKAN karena dia tidak becus membaca "
            "peta, tapi karena satuan skala yang dipakai berbeda dari yang "
            "dia pelajari. Itulah sebabnya tiga model (ConvNeXt Large, "
            "Base, dan Tiny) pada eksperimen awal akurasinya jatuh ke "
            "sekitar 48-53% (setara tebak koin) meski akurasi validasi "
            "saat training sempat terlihat baik (88-97%) -- bukan karena "
            "modelnya buruk, tapi karena \"bahasa angka\" yang dipakai saat "
            "diuji berbeda dari yang dipelajari saat dilatih.",
            "Normal",
        ),
        (
            "Satu model (ConvNeXt Small) kebetulan menunjukkan akurasi "
            "tinggi (98,6%) pada eksperimen awal -- TETAPI angka ini juga "
            "TIDAK BISA DIPERCAYA, karena diuji memakai data test yang "
            "sudah bocor 34,84% (Temuan 1 di atas). Akurasi tinggi yang "
            "kebetulan muncul di sini kemungkinan besar adalah hafalan "
            "terhadap gambar duplikat, bukan bukti model itu \"paling baik\" "
            "dibanding tiga model lainnya.",
            "Normal",
        ),
        ("Perbaikan yang Diterapkan pada Pipeline Final", "Heading 3"),
        (
            "Berdasarkan dua temuan di atas, pipeline final penelitian ini "
            "melakukan perbaikan berikut: (1) deduplikasi tingkat klaster -- "
            "setiap kelompok gambar identik/near-identik ditugaskan ke SATU "
            "split saja (train ATAU val ATAU test, tidak pernah lebih dari "
            "satu), sehingga kebocoran menjadi NOL secara konstruksi; "
            "(2) satu sumber kebenaran preprocessing (src/fracture/data.py), "
            "dipakai identik di training maupun saat model dipakai melayani "
            "prediksi di server, menghilangkan risiko preprocessing "
            "train≠test; (3) threshold keputusan klasifikasi dicari "
            "dari data VALIDASI (bukan data uji) memakai Youden's J index -- "
            "mencari threshold di data uji adalah bentuk kebocoran data lain "
            "yang juga ditemukan pada eksperimen awal (analoginya: "
            "menentukan nilai batas kelulusan SETELAH melihat hasil ujian "
            "siswa, supaya persentase kelulusan terlihat bagus -- itu bukan "
            "evaluasi yang jujur, karena standar kelulusan seharusnya "
            "ditentukan SEBELUM tahu hasilnya); dan (4) setiap metrik "
            "dilaporkan dengan interval kepercayaan 95% dari 2.000 kali "
            "resampling bootstrap, supaya perbedaan performa antar model "
            "hanya diklaim signifikan kalau rentang interval kepercayaannya "
            "TIDAK saling tumpang tindih -- bukan sekadar dibandingkan dari "
            "satu angka titik saja.",
            "Normal",
        ),
        (
            "Dengan perbaikan ini, seluruh angka yang dilaporkan pada "
            "subbab 4.6 \"Evaluasi Performa\" dan seterusnya di laporan ini "
            "adalah hasil dari pipeline yang sudah diperbaiki -- BUKAN "
            "kelanjutan dari eksperimen awal yang cacat metodologis di "
            "atas.",
            "Normal",
        ),
    ])

    print("Pass 10 selesai: disclaimer + section audit dataset baru disisipkan.")

    doc.save(DOCX_PATH)
    print(f"\nTersimpan: {DOCX_PATH}")


if __name__ == "__main__":
    main()
